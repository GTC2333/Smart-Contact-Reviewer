"""
Word exporter for contract audit reports.
"""
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from typing import Dict, Any
import tempfile


class WordExporter:
    """Export audit results to Word format."""

    def export(self, audit_result: Dict[str, Any], output_path: str = None) -> str:
        """
        Export audit result to Word document.

        Args:
            audit_result: Audit result dictionary
            output_path: Output file path (optional, returns temp path if None)

        Returns:
            Path to exported Word file
        """
        if output_path is None:
            output_path = tempfile.mktemp(suffix=".docx")

        doc = Document()

        # Title
        title = doc.add_heading('合同审核报告', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Contract info
        contract_id = audit_result.get("contract_id", "N/A")
        metadata = audit_result.get("metadata", {})
        processed_at = metadata.get("processed_at", "N/A")
        clause_count = metadata.get("clause_count", 0)
        party_count = metadata.get("party_count", 0)

        doc.add_paragraph(f"合同ID: {contract_id}")
        doc.add_paragraph(f"审核时间: {processed_at}")
        doc.add_paragraph(f"条款数量: {clause_count}")
        doc.add_paragraph(f"当事方数量: {party_count}")
        doc.add_paragraph()

        # Risk statistics (severity is in Chinese: 高/中/低)
        annotations = audit_result.get("annotations", [])
        high = len([a for a in annotations if a.get("severity", "").strip() == "高"])
        medium = len([a for a in annotations if a.get("severity", "").strip() == "中"])
        low = len([a for a in annotations if a.get("severity", "").strip() == "低"])

        doc.add_heading('风险统计', level=1)
        table = doc.add_table(rows=5, cols=2)
        table.style = 'Table Grid'

        data = [
            ("风险等级", "数量"),
            ("高风险", str(high)),
            ("中风险", str(medium)),
            ("低风险", str(low)),
            ("总计", str(len(annotations)))
        ]

        for i, (label, value) in enumerate(data):
            table.rows[i].cells[0].text = label
            table.rows[i].cells[1].text = value

        doc.add_paragraph()

        # Risk details
        if annotations:
            doc.add_heading('风险详情', level=1)

            for i, anno in enumerate(annotations[:20], 1):
                severity = anno.get("severity", "Unknown").upper()
                issue_type = anno.get("issue_type", "Risk")
                description = anno.get("description", "N/A")
                recommendation = anno.get("recommendation", "")

                p = doc.add_paragraph()
                p.add_run(f"{i}. [{severity}] {issue_type}").bold = True

                doc.add_paragraph(f"问题描述: {description}")
                if recommendation:
                    doc.add_paragraph(f"建议: {recommendation}")
                doc.add_paragraph()

        # Corrections
        corrections = audit_result.get("corrections", [])
        if corrections:
            doc.add_heading('修改建议', level=1)

            for i, corr in enumerate(corrections[:10], 1):
                revision = corr.get("suggested_revision", "N/A")
                note = corr.get("note", "")

                p = doc.add_paragraph()
                p.add_run(f"{i}. 修改建议: ").bold = True
                p.add_run(revision)

                if note:
                    doc.add_paragraph(f"备注: {note}")
                doc.add_paragraph()

        # Save
        doc.save(output_path)
        return output_path
